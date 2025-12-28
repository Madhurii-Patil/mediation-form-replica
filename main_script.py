import docx
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_final_custom_height_replica():
    """
    Main function to generate the exact replica of the PDF form in MS Word.
    """
    # 1. Initialize Document
    doc = Document()

    # 2. Setup Page Size (A4) and Margins (0.3")
    # We access the first section of the document to set dimensions.
    section = doc.sections[0]
    section.page_height = Mm(297) 
    section.page_width = Mm(210)  
    
    # Setting narrow margins to maximize printable area as per the PDF
    margin_size = Inches(0.3)
    section.top_margin = margin_size
    section.bottom_margin = margin_size
    section.left_margin = margin_size
    section.right_margin = margin_size

    # --- 3. Helper Functions (To avoid repeating code) ---

    def set_font(run, font_name='Times New Roman', font_size=10.5, bold=False, underline=False, color=None):
        """Applies specific font styling to a text run."""
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.underline = underline
        if color:
            run.font.color.rgb = color

    def add_centered_line(text, bold=True, size=11):
        """Adds a centered paragraph for the document headers."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Remove default spacing to keep lines tight
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        set_font(run, bold=bold, font_size=size)
        return p

    def set_cell_borders(cell):
        """
        Directly manipulates the underlying XML of a table cell to force 
        'Single' black borders. This is necessary because python-docx 
        defaults sometimes don't render all borders visible.
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for border_name in ['top', 'left', 'bottom', 'right']:
            tag = 'w:{}'.format(border_name)
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single') # Single line border
            element.set(qn('w:sz'), '4')       # Border size
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto') # Black color
            
            existing = tcBorders.find(qn(tag))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(element)

    def fill_cell(cell, text=None, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, underline=False, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER):
        """
        Fills a specific table cell with text and sets alignment/styling.
        Default vertical alignment is set to CENTER (Middle).
        """
        cell.text = "" # Clear default empty paragraph
        
        # Set Vertical Alignment (Top, Center, or Bottom)
        cell.vertical_alignment = vertical_align
            
        p = cell.add_paragraph()
        p.alignment = align
        
        # Set compact spacing for text inside cells
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        if text:
            run = p.add_run(text)
            set_font(run, bold=bold, color=color, underline=underline)
        
        # Apply the custom borders
        set_cell_borders(cell)
        return p

    # --- 4. Content Construction ---

    # Add Header Lines
    add_centered_line("FORM 'A'", bold=True, size=12)
    add_centered_line("MEDIATION APPLICATION FORM", bold=True, size=12)
    add_centered_line("[REFER RULE 3(1)]", bold=True, size=11)
    add_centered_line("Mumbai District Legal Services Authority", bold=False, size=12)
    p = add_centered_line("City Civil Court, Mumbai", bold=False, size=12)
    p.paragraph_format.space_after = Pt(10) # Add space before table starts

    # Create the Main Table
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER # Center table on page
    table.autofit = False 
    table.allow_autofit = False 
    
    # Define exact column widths to match the PDF
    col_widths = [Inches(0.37), Inches(1.2), Inches(5.68)]

    def add_row(height_in_inches=None):
        """Adds a row and enforces a specific height if provided."""
        row = table.add_row()
        if height_in_inches:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(height_in_inches)
        else:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            
        for i, width in enumerate(col_widths):
            row.cells[i].width = width
        return row

    # --- Row-by-Row Data Entry ---

    # ROW 1: DETAILS OF PARTIES
    row = add_row(0.44)
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    fill_cell(row.cells[0], "DETAILS OF PARTIES:", bold=True)

    # ROW 2: Name | Client Name
    row = add_row(0.45)
    fill_cell(row.cells[0], "1", align=WD_ALIGN_PARAGRAPH.CENTER, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    
    cell = row.cells[1]
    cell.text = "" 
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("Name of\nApplicant")
    set_font(r, bold=True)
    set_cell_borders(cell)
    
    fill_cell(row.cells[2], "{{client_name}}", bold=True, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.TOP)

    # ROW 3: Address Header
    row = add_row(0.37)
    set_cell_borders(row.cells[0]) 
    row.cells[1].merge(row.cells[2])
    fill_cell(row.cells[1], "Address and contact details of Applicant", bold=True)

    # ROW 4: Address Content
    row = add_row(1.12)
    fill_cell(row.cells[0], "1", align=WD_ALIGN_PARAGRAPH.CENTER, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    fill_cell(row.cells[1], "Address", bold=True, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    
    # Complex Cell (Registered & Correspondence Address)
    cell = row.cells[2]
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Registered Label
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("REGISTERED ADDRESS:")
    set_font(r, bold=True)
    
    # Variable
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("{{branch_address}}")
    set_font(r, bold=False)
    
    # Spacer
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(" ") 
    set_font(r, bold=False)
    
    # Correspondence Label
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("CORRESPONDENCE BRANCH ADDRESS:")
    set_font(r, bold=True)

    # Variable
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("{{branch_address}}")
    set_font(r, bold=False)
    
    set_cell_borders(cell)

    # ROW 5: Telephone
    row = add_row(0.37)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Telephone No.", bold=True)
    fill_cell(row.cells[2], "{{mobile}}", bold=True)

    # ROW 6: Mobile
    row = add_row(0.37)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Mobile No.", bold=True)
    set_cell_borders(row.cells[2])

    # ROW 7: Email
    row = add_row(0.37)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Email ID", bold=True)
    fill_cell(row.cells[2], "info@kslegal.co.in", color=RGBColor(0, 0, 255), underline=True)

    # ROW 8: Opposite Party Header
    row = add_row(0.37)
    fill_cell(row.cells[0], "2", align=WD_ALIGN_PARAGRAPH.CENTER, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    row.cells[1].merge(row.cells[2])
    fill_cell(row.cells[1], "Name, Address and Contact details of Opposite Party:", bold=True)

    # ROW 9: Defendant Header
    row = add_row(0.37)
    set_cell_borders(row.cells[0])
    row.cells[1].merge(row.cells[2])
    fill_cell(row.cells[1], "Address and contact details of Defendant/s", bold=True)

    # ROW 10: Name
    row = add_row(0.37)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Name", bold=True, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.TOP)
    fill_cell(row.cells[2], "{{customer_name}}", bold=True, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.TOP)

    # ROW 11: Address Jinja Block
    row = add_row(1.58)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Address", bold=True, vertical_align=WD_CELL_VERTICAL_ALIGNMENT.CENTER)
    
    cell = row.cells[2]
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Jinja Logic construction
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run("REGISTERED ADDRESS:")
    set_font(r, bold=True)
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("{% if address1 and address1 != \"\" %}{{address1}} {% else %} ________________ {%\n endif %}")
    set_font(r, bold=False)
    
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(" ") 
    set_font(r, bold=False)
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("CORRESPONDENCE ADDRESS:")
    set_font(r, bold=True)
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("{% if address1 and address1 != \"\" %}{{address1}} {% else %} ________________ {%\n endif %}")
    set_font(r, bold=False)
    
    set_cell_borders(cell)

    # ROW 12-14: Contacts
    row = add_row(0.44)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Telephone No.", bold=True)
    set_cell_borders(row.cells[2])

    row = add_row(0.44)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Mobile No.", bold=True)
    set_cell_borders(row.cells[2])

    row = add_row(0.44)
    set_cell_borders(row.cells[0])
    fill_cell(row.cells[1], "Email ID", bold=True)
    set_cell_borders(row.cells[2])

    # ROW 15: Details of Dispute
    row = add_row(0.44)
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    fill_cell(row.cells[0], "DETAILS OF DISPUTE:", bold=True)

    # ROW 16: Footer Title
    row = add_row(0.44)
    row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    cell = row.cells[0]
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    
    r = p.add_run("THE COMM. COURTS (PRE-INSTITUTION.........SETTLEMENT) RULES,2018")
    set_font(r, bold=True, underline=True)
    set_cell_borders(cell)

    # ROW 17: Footer Text
    row = add_row(0.22)
    set_cell_borders(row.cells[0])
    row.cells[1].merge(row.cells[2])
    fill_cell(row.cells[1], "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):", bold=True)

    # 5. Save the file
    # Updated to a more professional filename
    file_name = "Form_A_Mediation_Replica.docx"
    doc.save(file_name)
    print(f"File created successfully: {file_name}")

if __name__ == "__main__":
    create_final_custom_height_replica()